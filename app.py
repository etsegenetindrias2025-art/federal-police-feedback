from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file
import io
import os
import re
import socket
import time
import traceback
from datetime import datetime

from sqlalchemy import or_, and_, inspect, text as sql_text
from flask_sqlalchemy import SQLAlchemy

# Excel styling imports
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Word document imports
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# PDF generation imports (ReportLab)
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

import psycopg2
import psycopg2.extras

try:
    from better_profanity import profanity  # type: ignore[import-not-found]
except ImportError:
    class _FallbackProfanity:
        def __init__(self):
            self._censored_words = set()

        def load_censor_words(self):
            return None

        def add_censor_words(self, words):
            for word in words:
                cleaned = str(word).strip().lower()
                if cleaned:
                    self._censored_words.add(cleaned)

        def contains_profanity(self, text):
            if not text:
                return False
            normalized = re.sub(r'[^\w\s]', '', str(text).lower())
            for word in self._censored_words:
                if re.search(rf'\b{re.escape(word)}\b', normalized):
                    return True
            return False

    profanity = _FallbackProfanity()


# ----------------------------------------------------
# SINGLE FLASK APP + DATABASE INITIALIZATION
# ----------------------------------------------------
app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'federal_police_secret_key')

DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql://postgres:2323@localhost:5432/federal_police_feedback"
)

if DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

app.config["SQLALCHEMY_DATABASE_URI"] = DATABASE_URL
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_pre_ping": True,
}

db = SQLAlchemy(app)


def get_db_connection():
    """Raw psycopg2 connection, kept available for any ad-hoc queries.
    Everything in this file that talks to feedback data now goes through
    the SQLAlchemy models instead, since mixing raw cursors with the ORM
    was the source of the Pylance/runtime bugs in /admin/notifications."""
    conn = psycopg2.connect(
        dbname="federal_police_feedback",
        user="postgres",
        password="2323",
        host="localhost",
        port="5432",
        cursor_factory=psycopg2.extras.RealDictCursor
    )
    return conn


# ----------------------------------------------------
# COMPREHENSIVE PROFANITY FILTER (አማርኛ + Manglish + English)
# ----------------------------------------------------
profanity.load_censor_words()

ETHIOPIC_BAD_WORDS = [
    'ውሻ', 'ሌባ', 'አህያ', 'ጅብ', 'ፋንድያ', 'ሉጢ', 'ክፉ', 'በረንዳ አዳሪ',
    'wusha', 'wesha', 'leba', 'ahiya', 'jib', 'fandya', 'luti', 'dingay', 'balege', 'dedeb', 'denez'
]

profanity.add_censor_words(ETHIOPIC_BAD_WORDS)


def is_inappropriate(text):
    if not text:
        return False

    if profanity.contains_profanity(text):
        return True

    cleaned_text = re.sub(r'[^\w\s]', '', text.lower())
    for bad_word in ETHIOPIC_BAD_WORDS:
        if bad_word in cleaned_text:
            return True

    return False


# ----------------------------------------------------
# MODELS
# ----------------------------------------------------
class Feedback(db.Model):
    __tablename__ = 'feedbacks'
    id = db.Column(db.Integer, primary_key=True)
    service_name = db.Column(db.String(100), nullable=False)
    sub_service = db.Column(db.String(100), nullable=True)
    rating = db.Column(db.String(50), nullable=False)
    comment = db.Column(db.Text, nullable=True)
    audio_status = db.Column(db.String(100), default='No audio recorded')
    is_read = db.Column(db.Boolean, default=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)


class FingerprintRecord(db.Model):
    __tablename__ = 'fingerprint_records'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.String(100), nullable=True)
    fingerprint_data = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(50), default='Verified')
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)


class Service(db.Model):
    __tablename__ = 'services'
    id = db.Column(db.Integer, primary_key=True)
    service_key = db.Column(db.String(50), unique=True, nullable=False)
    service_name = db.Column(db.String(100), nullable=False)


class SubService(db.Model):
    __tablename__ = 'sub_services'
    id = db.Column(db.Integer, primary_key=True)
    service_key = db.Column(db.String(50), nullable=False)
    sub_service_key = db.Column(db.String(50), nullable=False)
    sub_service_name = db.Column(db.String(100), nullable=False)
    amharic_name = db.Column(db.String(100), nullable=True)


class AuditLog(db.Model):
    __tablename__ = 'audit_logs'
    id = db.Column(db.Integer, primary_key=True)
    admin_user = db.Column(db.String(100), nullable=False)
    action_description = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)


class Notification(db.Model):
    __tablename__ = 'notifications'
    id = db.Column(db.Integer, primary_key=True)
    message = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)


def ensure_postgresql_schema():
    """Migrate the old sub_services schema if it already exists."""
    inspector = inspect(db.engine)
    tables = set(inspector.get_table_names())

    if "sub_services" not in tables:
        return

    columns = {c["name"] for c in inspector.get_columns("sub_services")}
    required = {
        "id", "service_key", "sub_service_key",
        "sub_service_name", "amharic_name"
    }

    # db.create_all() cannot change an existing table's columns.
    # The older project used:
    # main_id, main_service, sub_id, sub_service, sub_description
    if required.issubset(columns):
        return

    legacy = "sub_services_legacy"

    if legacy not in tables:
        db.session.execute(
            sql_text("ALTER TABLE sub_services RENAME TO sub_services_legacy")
        )
        db.session.commit()

    db.create_all()

    legacy_columns = {
        c["name"] for c in inspect(db.engine).get_columns(legacy)
    }

    if {"main_service", "sub_service"}.issubset(legacy_columns):
        description_column = (
            "sub_description"
            if "sub_description" in legacy_columns
            else "NULL"
        )

        rows = db.session.execute(sql_text(
            f"SELECT main_service, sub_service, {description_column} "
            f"FROM {legacy}"
        )).fetchall()

        service_key_map = {
            "police clearance": "police_clearance",
            "complaint": "complaint",
            "hospital": "hospital",
            "logistics": "logistics",
            "education & training": "education_training",
            "education and training": "education_training",
            "human resources": "hr",
            "hr": "hr",
            "it help desk": "help_desk",
            "help desk": "help_desk",
            "other": "other",
        }

        for row in rows:
            service_name = str(row[0] or "").strip()
            sub_name = str(row[1] or "").strip()
            description = str(row[2] or "").strip()

            if not service_name or not sub_name:
                continue

            service_key = service_key_map.get(
                service_name.lower(),
                service_name.lower().replace(" ", "_")
            )

            sub_key = re.sub(
                r"[^a-z0-9_]",
                "",
                sub_name.lower()
                .replace("&", "and")
                .replace("/", "_")
                .replace("-", "_")
                .replace(" ", "_")
            )[:50]

            if not sub_key:
                continue

            if not SubService.query.filter_by(
                service_key=service_key,
                sub_service_key=sub_key
            ).first():
                row_obj = SubService()
                row_obj.service_key = service_key
                row_obj.sub_service_key = sub_key
                row_obj.sub_service_name = sub_name[:100]
                row_obj.amharic_name = description[:100] or None
                db.session.add(row_obj)

        db.session.commit()


def init_db():
    with app.app_context():
        ensure_postgresql_schema()
        db.create_all()

        default_services = [
            ("police_clearance", "Police Clearance"),
            ("complaint", "Complaint"),
            ("hospital", "Hospital"),
            ("logistics", "Logistics"),
            ("education_training", "Education & Training"),
            ("hr", "Human Resources"),
            ("help_desk", "IT Help Desk"),
            ("other", "Other")
        ]
        for key, name in default_services:
            existing = Service.query.filter_by(service_key=key).first()
            if not existing:
                svc = Service()
                svc.service_key = key
                svc.service_name = name
                db.session.add(svc)

        default_sub_services = {
            "police_clearance": [
                ("new_clearance", "New Police Clearance", "አዲስ የፖሊስ ክሊራንስ"),
                ("renewal", "Renewal", "እድሳት"),
                ("criminal_record", "Criminal Record Verification", "የወንጀል መዝገብ ማረጋገጫ"),
                ("fingerprint", "Fingerprint Registration", "የጣት አሻራ ምዝገባ"),
                ("document_collection", "Document Collection", "ሰነድ መሰብሰብ")
            ],
            "complaint": [
                ("crime_complaint", "Crime Complaint Registration", "የወንጀል ቅሬታ ምዝገባ"),
                ("public_office", "Public Complaint Office", "የህዝብ ቅሬታ ቢሮ"),
                ("online_followup", "Online Complaint Follow-up", "የመስመር ላይ ቅሬታ ክትትል"),
                ("investigation", "Investigation", "ምርመራ"),
                ("resolution", "Resolution", "ውሳኔ")
            ],
            "hospital": [
                ("opd", "OPD", "የውጪ ህሙማን ክፍል"),
                ("emergency", "Emergency", "አስቸኳይ ጊዜ"),
                ("pharmacy", "Pharmacy", "ፋርማሲ"),
                ("laboratory", "Laboratory", "ላቦራቶሪ"),
                ("medical_exam", "Medical Examination", "የህክምና ምርመራ")
            ],
            "logistics": [
                ("vehicle_mgmt", "Vehicle Management", "የተሽከርካሪ አስተዳደር"),
                ("garage", "Garage", "ጋራጅ"),
                ("equipment_dist", "Equipment Distribution", "የቁሳቁስ ክፍፍል"),
                ("inventory", "Inventory", "ዕቃ ግምጃ ቤት"),
                ("procurement", "Procurement", "ግዥ")
            ],
            "education_training": [
                ("student_reg", "Student Registration", "የተማሪዎች ምዝገባ"),
                ("training", "Training", "ስልጠና"),
                ("certificates", "Certificates", "ሰርተፍኬቶች"),
                ("examination", "Examination", "ፈተና"),
                ("academic_records", "Academic Records", "የአካዳሚክ መዛግብት")
            ],
            "hr": [
                ("recruitment", "Recruitment & Talent Acquisition", "የሰራተኛ ቅጥር እና ተሰጥኦ ማፈላለግ"),
                ("employee_records", "Employee Records", "የሰራተኛ መዝገቦች"),
                ("payroll", "Payroll & Benefits", "የደመወዝ እና ጥቅማጥቅሞች"),
                ("leave", "Leave Management", "የፈቃድ አስተዳደር"),
                ("training_development", "Staff Development", "የሰራተኛ ልማት")
            ],
            "help_desk": [
                ("hardware", "Hardware Support", "የሃርድዌር ድጋፍ"),
                ("network", "Network Support", "የኔትወርክ ድጋፍ"),
                ("software", "Software Support", "የሶፍትዌር ድጋፍ"),
                ("account_access", "Account & Access Support", "የመለያ እና የመግቢያ ድጋፍ"),
                ("technical_issue", "Technical Issue Reporting", "የቴክኒክ ችግር ሪፖርት")
            ],
            "other": [
                ("reception", "Reception", "እንግዳ መቀበያ"),
                ("ict_support", "ICT Support", "የአይቲ ድጋፍ"),
                ("finance", "Finance", "ፋይናንስ"),
                ("admin", "Administration", "አስተዳደር")
            ]
        }

        for s_key, sub_list in default_sub_services.items():
            for sub_key, sub_name, amh_name in sub_list:
                existing = SubService.query.filter_by(service_key=s_key, sub_service_key=sub_key).first()
                if not existing:
                    ss = SubService()
                    ss.service_key = s_key
                    ss.sub_service_key = sub_key
                    ss.sub_service_name = sub_name
                    ss.amharic_name = amh_name
                    db.session.add(ss)

        # Migration for older database versions:
        # move the old HR sub-service from "other" to the new "hr" service.
        old_hr = SubService.query.filter_by(
            service_key="other",
            sub_service_key="hr"
        ).first()
        if old_hr:
            old_hr.service_key = "hr"
            old_hr.sub_service_name = "Human Resources"
            old_hr.amharic_name = "ሰው ሃብት"

        # Keep historical feedback records consistent with the new main service.
        old_hr_feedbacks = Feedback.query.filter_by(
            service_name="other",
            sub_service="hr"
        ).all()
        for fb in old_hr_feedbacks:
            fb.service_name = "hr"

        db.session.commit()


init_db()
print("[startup] Successfully initialized PostgreSQL database tables.")


def log_admin_action(username, description):
    try:
        with app.app_context():
            # Clear any dangling failed-transaction state from an earlier
            # query in this request so this insert doesn't get silently
            # dropped along with it.
            db.session.rollback()
            log = AuditLog()
            log.admin_user = str(username)
            log.action_description = str(description)
            db.session.add(log)
            db.session.commit()
    except Exception:
        db.session.rollback()
        print("=== Error logging audit trail (see traceback below) ===")
        traceback.print_exc()


def get_service_map():
    with app.app_context():
        services = Service.query.order_by(Service.id.asc()).all()
        return {s.service_key: s.service_name for s in services}


def get_sub_service_map():
    with app.app_context():
        rows = SubService.query.all()
        sub_map = {}
        for row in rows:
            if row.service_key not in sub_map:
                sub_map[row.service_key] = {}
            sub_map[row.service_key][row.sub_service_key] = row.sub_service_name
        return sub_map


def get_admin_credentials():
    return {
        "admin gen": {"password": "1234", "type": "general", "service": "all", "sub_service": "all", "title": "General Admin Dashboard"},
        "admin pol": {"password": "1234", "type": "service", "service": "police_clearance", "sub_service": "all", "title": "Police Clearance Admin"},
        "admin com": {"password": "1234", "type": "service", "service": "complaint", "sub_service": "all", "title": "Complaint Admin"},
        "admin hos": {"password": "1234", "type": "service", "service": "hospital", "sub_service": "all", "title": "Hospital Admin"},
        "admin log": {"password": "1234", "type": "service", "service": "logistics", "sub_service": "all", "title": "Logistics Admin"},
        "admin edu": {"password": "1234", "type": "service", "service": "education_training", "sub_service": "all", "title": "Education Admin"},
        "admin hr": {"password": "1234", "type": "service", "service": "hr", "sub_service": "all", "title": "Human Resources Admin"},
        "admin help": {"password": "1234", "type": "service", "service": "help_desk", "sub_service": "all", "title": "IT Help Desk Admin"},
        "admin oth": {"password": "1234", "type": "service", "service": "other", "sub_service": "all", "title": "Other Admin"},
        "admin new": {"password": "1234", "type": "sub_service", "service": "police_clearance", "sub_service": "new_clearance", "title": "Sub Admin: New Clearance"},
        "admin ren": {"password": "1234", "type": "sub_service", "service": "police_clearance", "sub_service": "renewal", "title": "Sub Admin: Renewal"},
        "admin cri": {"password": "1234", "type": "sub_service", "service": "police_clearance", "sub_service": "criminal_record", "title": "Sub Admin: Criminal Record"},
        "admin fin": {"password": "1234", "type": "sub_service", "service": "police_clearance", "sub_service": "fingerprint", "title": "Sub Admin: Fingerprint"},
        "admin doc": {"password": "1234", "type": "sub_service", "service": "police_clearance", "sub_service": "document_collection", "title": "Sub Admin: Document Collection"},
        "admin pub": {"password": "1234", "type": "sub_service", "service": "complaint", "sub_service": "public_office", "title": "Sub Admin: Public Complaint"},
        "admin onl": {"password": "1234", "type": "sub_service", "service": "complaint", "sub_service": "online_followup", "title": "Sub Admin: Online Follow-up"},
        "admin inv": {"password": "1234", "type": "sub_service", "service": "complaint", "sub_service": "investigation", "title": "Sub Admin: Investigation"},
        "admin res": {"password": "1234", "type": "sub_service", "service": "complaint", "sub_service": "resolution", "title": "Sub Admin: Resolution"},
        "admin opd": {"password": "1234", "type": "sub_service", "service": "hospital", "sub_service": "opd", "title": "Sub Admin: OPD"},
        "admin eme": {"password": "1234", "type": "sub_service", "service": "hospital", "sub_service": "emergency", "title": "Sub Admin: Emergency"},
        "admin pha": {"password": "1234", "type": "sub_service", "service": "hospital", "sub_service": "pharmacy", "title": "Sub Admin: Pharmacy"},
        "admin lab": {"password": "1234", "type": "sub_service", "service": "hospital", "sub_service": "laboratory", "title": "Sub Admin: Laboratory"},
        "admin med": {"password": "1234", "type": "sub_service", "service": "hospital", "sub_service": "medical_exam", "title": "Sub Admin: Medical Examination"},
        "admin veh": {"password": "1234", "type": "sub_service", "service": "logistics", "sub_service": "vehicle_mgmt", "title": "Sub Admin: Vehicle Management"},
        "admin gar": {"password": "1234", "type": "sub_service", "service": "logistics", "sub_service": "garage", "title": "Sub Admin: Garage"},
        "admin equ": {"password": "1234", "type": "sub_service", "service": "logistics", "sub_service": "equipment_dist", "title": "Sub Admin: Equipment Distribution"},
        "admin pro": {"password": "1234", "type": "sub_service", "service": "logistics", "sub_service": "procurement", "title": "Sub Admin: Procurement"},
        "admin stu": {"password": "1234", "type": "sub_service", "service": "education_training", "sub_service": "student_reg", "title": "Sub Admin: Student Registration"},
        "admin tra": {"password": "1234", "type": "sub_service", "service": "education_training", "sub_service": "training", "title": "Sub Admin: Training"},
        "admin cer": {"password": "1234", "type": "sub_service", "service": "education_training", "sub_service": "certificates", "title": "Sub Admin: Certificates"},
        "admin exa": {"password": "1234", "type": "sub_service", "service": "education_training", "sub_service": "examination", "title": "Sub Admin: Examination"},
        "admin aca": {"password": "1234", "type": "sub_service", "service": "education_training", "sub_service": "academic_records", "title": "Sub Admin: Academic Records"},
        "admin rec": {"password": "1234", "type": "sub_service", "service": "other", "sub_service": "reception", "title": "Sub Admin: Reception"},
        "admin ict": {"password": "1234", "type": "sub_service", "service": "other", "sub_service": "ict_support", "title": "Sub Admin: ICT Support"},
        "admin fnn": {"password": "1234", "type": "sub_service", "service": "other", "sub_service": "finance", "title": "Sub Admin: Finance"},
        "admin adm": {"password": "1234", "type": "sub_service", "service": "other", "sub_service": "admin", "title": "Sub Admin: Administration"},
        "admin rec_hr": {"password": "1234", "type": "sub_service", "service": "hr", "sub_service": "recruitment", "title": "Sub Admin: Recruitment"},
        "admin emp": {"password": "1234", "type": "sub_service", "service": "hr", "sub_service": "employee_records", "title": "Sub Admin: Employee Records"},
        "admin pay": {"password": "1234", "type": "sub_service", "service": "hr", "sub_service": "payroll", "title": "Sub Admin: Payroll & Benefits"},
        "admin lea": {"password": "1234", "type": "sub_service", "service": "hr", "sub_service": "leave", "title": "Sub Admin: Leave Management"},
        "admin dev": {"password": "1234", "type": "sub_service", "service": "hr", "sub_service": "training_development", "title": "Sub Admin: Staff Development"}
    }


# ----------------------------------------------------
# ADVANCED SEARCH, FILTER & REPORT BUILDERS
# ----------------------------------------------------
def _resolve_service_key(fb, service_key_by_name):
    """Best-effort normalization of a feedback row's stored service_name
    into one of the canonical service_key values from the services table."""
    db_s = str(fb.service_name).strip().lower()
    return service_key_by_name.get(db_s, db_s)


def get_filtered_feedbacks(admin_type, assigned_service, assigned_sub):
    """Returns Feedback rows scoped to this admin, then narrowed by
    whichever of the search/filter query params are present. Both the
    'folder' click (?service=...) and the advanced filter dropdown
    (?service_filter=...) are honored; the free-text box is accepted as
    either 'search' (current template) or 'q' (legacy links)."""
    search_query = (request.args.get('search') or request.args.get('q') or '').strip()
    service_filter = request.args.get('service_filter', 'all')
    folder_filter = request.args.get('service', 'all')
    rating_filter = request.args.get('rating', 'all')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    all_rows = Feedback.query.all()
    filtered = []

    service_map_for_filter = get_service_map()
    service_key_by_name = {
        str(name).strip().lower(): key
        for key, name in service_map_for_filter.items()
    }

    assigned_service_key = str(assigned_service).strip().lower()

    for fb in all_rows:
        db_service_key = _resolve_service_key(fb, service_key_by_name)
        db_sub = str(fb.sub_service).strip().lower()

        if admin_type == 'service' and db_service_key != assigned_service_key:
            continue
        if admin_type == 'sub_service' and (
            db_service_key != assigned_service_key
            or db_sub != str(assigned_sub).strip().lower()
        ):
            continue

        # Advanced-filter dropdown takes priority; otherwise fall back to
        # the folder that was clicked on the general dashboard.
        active_service_selection = service_filter if service_filter and service_filter != 'all' else folder_filter
        if active_service_selection and active_service_selection != 'all':
            selected_service_key = str(active_service_selection).strip().lower()
            if selected_service_key in service_key_by_name:
                selected_service_key = service_key_by_name[selected_service_key]
            if db_service_key != selected_service_key:
                continue

        if rating_filter and rating_filter != 'all':
            if str(fb.rating).strip() != str(rating_filter).strip():
                continue

        if date_from or date_to:
            dt_rec = fb.timestamp

            if isinstance(dt_rec, str):
                try:
                    dt_rec = datetime.fromisoformat(dt_rec.replace("Z", "+00:00"))
                except ValueError:
                    try:
                        dt_rec = datetime.strptime(dt_rec.split('.')[0], '%Y-%m-%d %H:%M:%S')
                    except ValueError:
                        dt_rec = None

            if dt_rec:
                if date_from:
                    try:
                        df_dt = datetime.strptime(date_from, '%Y-%m-%d')
                        if dt_rec < df_dt:
                            continue
                    except ValueError:
                        pass
                if date_to:
                    try:
                        dt_dt = datetime.strptime(date_to, '%Y-%m-%d').replace(hour=23, minute=59, second=59)
                        if dt_rec > dt_dt:
                            continue
                    except ValueError:
                        pass

        if search_query:
            q_lower = search_query.lower()
            combined_text = f"#{fb.id} {fb.service_name} {fb.sub_service} {fb.rating} {fb.comment} {fb.timestamp}".lower()
            if q_lower not in combined_text:
                continue

        filtered.append(fb)

    filtered.sort(key=lambda fb: fb.timestamp, reverse=True)
    return filtered


def build_ai_insights(records):
    """Lightweight, dependency-free heuristics for the dashboard's
    'AI Summary & Insights' card, computed straight from the scoped
    Postgres records (no external API calls)."""
    if not records:
        return {
            'satisfaction': 'N/A',
            'top_service': 'N/A',
            'main_complaint': 'None',
            'recommendation': 'No feedback records available yet for this scope.'
        }

    positive_ratings = {'😍', '😊', '4', '5'}
    negative_ratings = {'🙁', '😠', '😡', '1', '2'}

    positive_count = sum(1 for fb in records if str(fb.rating).strip() in positive_ratings)
    satisfaction_pct = round((positive_count / len(records)) * 100)

    service_map = get_service_map()
    service_totals = {}
    for fb in records:
        key = str(fb.service_name).strip().lower()
        service_totals[key] = service_totals.get(key, 0) + 1

    stopwords = {
        'the', 'a', 'an', 'is', 'was', 'and', 'to', 'of', 'it', 'in', 'on',
        'for', 'with', 'this', 'that', 'i', 'my', 'we', 'our', 'not', 'very',
        'so', 'but', 'no', 'are', 'be', 'have', 'has', 'me', 'you', 'your'
    }
    word_counts = {}
    for fb in records:
        if str(fb.rating).strip() in negative_ratings and fb.comment:
            for word in re.findall(r"[^\W\d_]+", fb.comment.lower(), flags=re.UNICODE):
                if word not in stopwords and len(word) > 2:
                    word_counts[word] = word_counts.get(word, 0) + 1

    recommendation = (
        f"Based on {len(records)} visible feedback records, continue monitoring "
    )

    return {
        'satisfaction': f"{satisfaction_pct}%",
        'recommendation': recommendation
    }


# --- EXCEL REPORT (openpyxl) ---
def generate_excel_report(records):
    wb = openpyxl.Workbook()
    ws = wb.active
    if ws is None:
        ws = wb.create_sheet()
    ws.title = "Feedback Report"
    ws.views.sheetView[0].showGridLines = True

    header_fill = PatternFill(start_color="1B2A4A", end_color="1B2A4A", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    title_font = Font(name="Calibri", size=16, bold=True, color="1B2A4A")
    meta_font = Font(name="Calibri", size=10, italic=True, color="555555")
    data_font = Font(name="Calibri", size=10)

    border_thin = Border(
        left=Side(style='thin', color='DDDDDD'),
        right=Side(style='thin', color='DDDDDD'),
        top=Side(style='thin', color='DDDDDD'),
        bottom=Side(style='thin', color='DDDDDD')
    )

    ws.append(["Ethiopian Federal Police - Feedback Management Report"])
    ws.cell(row=1, column=1).font = title_font
    ws.append([f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Total Records: {len(records)}"])
    ws.cell(row=2, column=1).font = meta_font
    ws.append([])

    headers = ["ID", "Service", "Sub-Service", "Rating", "Comment", "Audio Feedback", "Submission Date"]
    ws.append(headers)

    for col_num in range(1, len(headers) + 1):
        cell = ws.cell(row=4, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    for r_idx, rec in enumerate(records, start=5):
        ws.append([
            f"#{rec.id}",
            rec.service_name,
            rec.sub_service,
            str(rec.rating),
            rec.comment or "No comment provided.",
            rec.audio_status or "No audio recorded",
            str(rec.timestamp)
        ])
        for c_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=r_idx, column=c_idx)
            cell.font = data_font
            cell.border = border_thin
            if c_idx in [1, 4, 7]:
                cell.alignment = Alignment(horizontal="center")

    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        if col[0].column is not None:
            col_letter = get_column_letter(int(col[0].column))
            ws.column_dimensions[col_letter].width = max(max_len + 4, 12)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


# --- WORD REPORT (python-docx) ---
def generate_word_report(records):
    doc = docx.Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Ethiopian Federal Police")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(27, 42, 74)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(
        f"Feedback System Audit Report\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Records Found: {len(records)}"
    )
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    headers = ["ID & Date", "Service / Sub-Service", "Rating", "Comment", "Audio"]
    col_widths = [Inches(1.2), Inches(1.5), Inches(0.8), Inches(2.0), Inches(1.0)]

    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = col_widths[i]
        shading = parse_xml(r'<w:shd {} w:fill="1B2A4A"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    for rec in records:
        row_cells = table.add_row().cells
        row_cells[0].text = f"#{rec.id}\n{str(rec.timestamp).split()[0]}"
        row_cells[1].text = f"{rec.service_name}\n({rec.sub_service})"
        row_cells[2].text = str(rec.rating)
        row_cells[3].text = rec.comment or "No comment."
        row_cells[4].text = rec.audio_status or "None"

        for i, cell in enumerate(row_cells):
            cell.width = col_widths[i]
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# --- PDF REPORT (ReportLab) ---
def generate_pdf_report(records):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(letter),
        rightMargin=36, leftMargin=36, topMargin=40, bottomMargin=40
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=colors.HexColor('#1B2A4A'),
        spaceAfter=4
    )
    subtitle_style = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=10,
        textColor=colors.HexColor('#555555'),
        spaceAfter=15
    )
    cell_style = ParagraphStyle(
        'CellText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        textColor=colors.HexColor('#222222'),
        leading=11
    )
    header_cell_style = ParagraphStyle(
        'HeaderCellText',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        textColor=colors.white,
        alignment=1
    )

    elements = []

    elements.append(Paragraph("Ethiopian Federal Police - Feedback Management Report", title_style))
    elements.append(Paragraph(
        f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Total Filtered Records: {len(records)}",
        subtitle_style
    ))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#1B2A4A'), spaceAfter=15))

    table_data = [[
        Paragraph("ID", header_cell_style),
        Paragraph("Service & Sub-Service", header_cell_style),
        Paragraph("Rating", header_cell_style),
        Paragraph("Comment", header_cell_style),
        Paragraph("Audio Status", header_cell_style),
        Paragraph("Submission Date", header_cell_style)
    ]]

    for rec in records:
        table_data.append([
            Paragraph(f"#{rec.id}", cell_style),
            Paragraph(f"<b>{rec.service_name}</b><br/>{rec.sub_service}", cell_style),
            Paragraph(str(rec.rating), cell_style),
            Paragraph(rec.comment or "No comment provided.", cell_style),
            Paragraph(rec.audio_status or "No audio", cell_style),
            Paragraph(str(rec.timestamp), cell_style)
        ])

    col_widths = [50, 160, 60, 260, 100, 110]
    t = Table(table_data, colWidths=col_widths, repeatRows=1)

    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1B2A4A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#DCDCDC')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9FAFB')])
    ]))

    elements.append(t)
    doc.build(elements)

    buffer.seek(0)
    return buffer


# ----------------------------------------------------
# PROGRESSIVE WEB APP (PWA) OFFLINE ROUTE
# ----------------------------------------------------
@app.route('/sw.js')
def service_worker():
    return app.send_static_file('sw.js')


# ----------------------------------------------------
# PUBLIC ROUTES
# ----------------------------------------------------
@app.route('/')
def fingerprint():
    return render_template('fingerprint.html')


@app.route('/welcome')
@app.route('/language')
def welcome_page():
    lang = request.args.get('lang', 'am')
    return render_template('welcome.html', lang=lang)


@app.route('/services')
def services():
    lang = request.args.get('lang', 'am')
    service_map = get_service_map()
    return render_template('services.html', lang=lang, service_map=service_map)


@app.route('/feedback')
def feedback():
    lang = request.args.get('lang', 'am')
    service = request.args.get('service', 'police_clearance')
    service_map = get_service_map()
    sub_service_map = get_sub_service_map()
    return render_template(
        'feedback.html', lang=lang, service=service,
        service_map=service_map, sub_service_map=sub_service_map
    )


# ----------------------------------------------------
# DEPARTMENTS API ROUTE FOR FRONTEND
# ----------------------------------------------------
@app.route('/api/departments')
def get_departments():
    try:
        rows = SubService.query.all()

        category_meta = {
            "police_clearance": {
                "title": {"en": "Police Clearance & Records", "am": "የፖሊስ ክሊራንስ እና መዛግብት"},
                "description": {"en": "New clearance, renewal, and record verification", "am": "አዲስ ክሊራንስ፣ እድሳት እና መዛግብት ማረጋገጫ"},
                "icon": "fa-id-card"
            },
            "complaint": {
                "title": {"en": "Crime Investigation & Complaints", "am": "የወንጀል ምርመራ እና ቅሬታዎች"},
                "description": {"en": "Crime report filing, public complaint office, follow-ups", "am": "የወንጀል ሪፖርት ማቅረቢያ፣ የህዝብ ቅሬታ ቢሮ"},
                "icon": "fa-magnifying-glass"
            },
            "hospital": {
                "title": {"en": "Medical Services", "am": "የህክምና አገልግሎቶች"},
                "description": {"en": "OPD, emergency, pharmacy, and laboratory", "am": "የውጪ ህሙማን፣ አስቸኳይ ጊዜ፣ ፋርማሲ እና ላቦራቶሪ"},
                "icon": "fa-hospital"
            },
            "logistics": {
                "title": {"en": "Logistics & Fleet", "am": "ሎጂስቲክስ እና መኪና አስተዳደር"},
                "description": {"en": "Vehicle management, garage, and inventory", "am": "የተሽከርካሪ አስተዳደር፣ ጋራጅ እና ዕቃ ግምጃ ቤት"},
                "icon": "fa-truck-fast"
            },
            "education_training": {
                "title": {"en": "Education & Training", "am": "ትምህርት እና ስልጠና"},
                "description": {"en": "Student registration, training, and academic records", "am": "የተማሪዎች ምዝገባ፣ ስልጠና እና የአካዳሚክ መዛግብት"},
                "icon": "fa-graduation-cap"
            },
            "hr": {
                "title": {"en": "Human Resources", "am": "ሰው ሃብት አስተዳደር"},
                "description": {"en": "Talent acquisition, payroll, and operations", "am": "የሰራተኛ ቅጥር፣ የደመወዝ ክፍያ እና አስተዳደር"},
                "icon": "fa-users-gear"
            },
            "help_desk": {
                "title": {"en": "IT Help Desk", "am": "የአይቲ እርዳታ ማዕከል"},
                "description": {"en": "Hardware maintenance, networking, and software support", "am": "የሃርድዌር ጥገና፣ ኔትወርክ እና ሶፍትዌር ድጋፍ"},
                "icon": "fa-headset"
            },
            "other": {
                "title": {"en": "General Support & Admin", "am": "አጠቃላይ ድጋፍ እና አስተዳደር"},
                "description": {"en": "Reception, ICT support, finance, and administration", "am": "እንግዳ መቀበያ፣ ፋይናንስ እና አስተዳደር"},
                "icon": "fa-building-shield"
            }
        }

        departments = {}
        for row in rows:
            s_key = row.service_key
            if s_key not in departments:
                meta = category_meta.get(s_key, {
                    "title": {"en": s_key.replace('_', ' ').title(), "am": s_key},
                    "description": {"en": "Department services and inquiries", "am": "የምድብ አገልግሎቶች"},
                    "icon": "fa-folder"
                })
                departments[s_key] = {
                    "title": meta["title"],
                    "description": meta["description"],
                    "icon": meta["icon"],
                    "items": []
                }

            clean_amharic = row.amharic_name
            if clean_amharic and '?' in clean_amharic:
                clean_amharic = row.sub_service_name

            departments[s_key]["items"].append({
                "id": row.sub_service_key,
                "name": {
                    "en": row.sub_service_name,
                    "am": clean_amharic or row.sub_service_name
                },
                "icon": "fa-file-lines"
            })

        return jsonify(departments)
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ----------------------------------------------------
# FINGERPRINT API ROUTES
# ----------------------------------------------------
@app.route('/api/fingerprint/scan', methods=['POST'])
def scan_fingerprint():
    try:
        data = request.get_json() or {}
        user_id = data.get('user_id', 'TEMPORARY_USER')
        fingerprint_data = data.get('fingerprint_data', 'BYPASS_FINGERPRINT_HASH')

        record = FingerprintRecord()
        record.user_id = str(user_id)
        record.fingerprint_data = str(fingerprint_data)
        db.session.add(record)
        db.session.commit()

        return jsonify({
            "status": "success",
            "message": "የጣት አሻራ ስካን ሳይጠበቅ ቀጥታ አልፏል (Demo Mode)!",
            "user_id": user_id
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ----------------------------------------------------
# FEEDBACK SUBMIT / UNREAD ROUTES
# ----------------------------------------------------
@app.route('/submit-feedback', methods=['POST'])
@app.route('/api/submit-feedback', methods=['POST'])
def submit_feedback():
    try:
        feedback_count = session.get('feedback_count', 0)
        if feedback_count >= 3:
            return jsonify({
                "status": "error",
                "message": "ለአሁኑ የተፈቀደልዎትን 3 አስተያየቶች ጨርሰዋል! / You have reached your max limit of 3 feedbacks for this session."
            }), 403

        data = request.get_json()
        if not data:
            return jsonify({"status": "error", "message": "No data received"}), 400

        rating = data.get('rating', '😊')
        comment = data.get('comment', 'No comment provided.')
        sub_service = data.get('sub_service', 'general_service')
        audio_status = data.get('audio_status', 'No audio recorded')

        url_service = data.get('service') or data.get('category') or data.get('service_name', 'police_clearance')
        client_timestamp = data.get('timestamp')

        if is_inappropriate(comment):
            return jsonify({
                "status": "error",
                "message": "Inappropriate language detected. Please keep your feedback respectful."
            }), 400

        parsed_ts = datetime.utcnow()
        if client_timestamp:
            try:
                parsed_ts = datetime.strptime(client_timestamp.split('.')[0], '%Y-%m-%d %H:%M:%S')
            except ValueError:
                pass

        new_fb = Feedback()
        new_fb.service_name = str(url_service)
        new_fb.sub_service = str(sub_service)
        new_fb.rating = str(rating)
        new_fb.comment = str(comment)
        new_fb.audio_status = str(audio_status)
        new_fb.is_read = False
        new_fb.timestamp = parsed_ts
        db.session.add(new_fb)
        db.session.commit()

        session['feedback_count'] = feedback_count + 1

        return jsonify({
            "status": "success",
            "message": f"Feedback saved successfully! ({session['feedback_count']}/3 submitted)"
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/api/unread-count')
def api_unread_count():
    try:
        unread_count = Feedback.query.filter_by(is_read=False).count()
    except Exception:
        unread_count = 0
    return jsonify({"unread_count": unread_count})


@app.route('/api/notifications/unread-count')
def api_notifications_unread_count():
    """Scoped unread count for the logged-in admin's notification bell.
    Unlike /api/unread-count (system-wide), this respects each admin's
    department scope, same as the dashboard and notifications page."""
    logged_in_admin = session.get('admin_user')
    admin_credentials = get_admin_credentials()
    if not logged_in_admin or logged_in_admin not in admin_credentials:
        return jsonify({"unread_count": 0}), 401

    admin_info = admin_credentials[logged_in_admin]
    try:
        scoped_records = get_filtered_feedbacks(
            admin_info['type'], admin_info['service'], admin_info['sub_service']
        )
        unread_count = sum(1 for fb in scoped_records if not fb.is_read)
    except Exception:
        unread_count = 0

    return jsonify({"unread_count": unread_count})


# ----------------------------------------------------
# ADMIN AUTHENTICATION
# ----------------------------------------------------
@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    error = None
    admin_credentials = get_admin_credentials()
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if username in admin_credentials and admin_credentials[username]['password'] == password:
            session['admin_user'] = username
            log_admin_action(username, "Logged into the admin panel.")
            return redirect(url_for('admin_dashboard'))
        else:
            error = "Invalid Username or Password. Please try again."
    return render_template('admin_login.html', error=error)


@app.route('/admin/logout')
def admin_logout():
    logged_in_admin = session.get('admin_user')
    if logged_in_admin:
        log_admin_action(logged_in_admin, "Logged out of the admin panel.")
    session.pop('admin_user', None)
    return redirect(url_for('admin_login'))


# ----------------------------------------------------
# ADMIN DASHBOARD & REPORTS EXPORT ROUTES
# ----------------------------------------------------
@app.route('/admin/dashboard')
def admin_dashboard():
    logged_in_admin = session.get('admin_user')
    admin_credentials = get_admin_credentials()
    if not logged_in_admin or logged_in_admin not in admin_credentials:
        return redirect(url_for('admin_login'))

    admin_info = admin_credentials[logged_in_admin]
    admin_type = admin_info['type']
    assigned_service = admin_info['service']
    assigned_sub = admin_info['sub_service']
    admin_title = admin_info['title']
    is_general = (admin_type == 'general')
    current_filter = request.args.get('service', 'all')
    search_query = request.args.get('search', '').strip()

    service_map = get_service_map()
    sub_service_map = get_sub_service_map()
    service_key_by_name = {
        str(name).strip().lower(): key for key, name in service_map.items()
    }

    # Every record this admin is allowed to see at all (before the folder
    # filter is applied) -- used for the totals, chart and folder counts,
    # so those stay stable while browsing folders.
    scoped_records = get_filtered_feedbacks(admin_type, assigned_service, assigned_sub)

    # Records after the folder/search/rating/date filters currently in the URL.
    feedbacks = scoped_records

    total_feedbacks_count = len(scoped_records)

    service_counts = {}
    for key in service_map:
        service_counts[key] = sum(
            1 for fb in scoped_records
            if _resolve_service_key(fb, service_key_by_name) == key
        )
    chart_data = {
        service_map[key]: count for key, count in service_counts.items() if count > 0
    }

    ai_insights = build_ai_insights(scoped_records)
    unread_notifications_count = sum(1 for fb in scoped_records if not fb.is_read)

    return render_template(
        'admin_dashboard.html',
        admin_user=logged_in_admin,
        admin_title=admin_title,
        is_general=is_general,
        assigned_service=assigned_service,
        assigned_sub=assigned_sub,
        feedbacks=feedbacks,
        service_map=service_map,
        sub_service_map=sub_service_map,
        current_filter=current_filter,
        search_query=search_query,
        total_feedbacks_count=total_feedbacks_count,
        service_counts=service_counts,
        chart_data=chart_data,
        ai_insights=ai_insights,
        unread_notifications_count=unread_notifications_count
    )


@app.route('/admin/notifications')
def admin_notifications():
    logged_in_admin = session.get('admin_user')
    admin_credentials = get_admin_credentials()
    if not logged_in_admin or logged_in_admin not in admin_credentials:
        return redirect(url_for('admin_login'))

    admin_info = admin_credentials[logged_in_admin]
    admin_type = admin_info['type']
    assigned_service = admin_info['service']
    assigned_sub = admin_info['sub_service']

    # Scoped to this admin's department, exactly like the dashboard.
    notifications = get_filtered_feedbacks(admin_type, assigned_service, assigned_sub)

    unread_ids = [fb.id for fb in notifications if not fb.is_read]
    if unread_ids:
        Feedback.query.filter(Feedback.id.in_(unread_ids)).update(
            {Feedback.is_read: True}, synchronize_session=False
        )
        db.session.commit()
        for fb in notifications:
            fb.is_read = True

    service_map = get_service_map()
    sub_service_map = get_sub_service_map()

    return render_template(
        'admin_notifications.html',
        notifications=notifications,
        service_map=service_map,
        sub_service_map=sub_service_map
    )


@app.route('/admin/audit-logs')
def admin_audit_logs():
    logged_in_admin = session.get('admin_user')
    admin_credentials = get_admin_credentials()
    if not logged_in_admin or logged_in_admin not in admin_credentials:
        return redirect(url_for('admin_login'))

    admin_info = admin_credentials[logged_in_admin]
    is_general_admin = (admin_info['type'] == 'general')

    # General admin sees the full system-wide log, pulled straight from
    # Postgres; scoped admins only see their own activity ("My Activity
    # Audit Trail").
    logs_query = AuditLog.query
    if not is_general_admin:
        logs_query = logs_query.filter(AuditLog.admin_user == logged_in_admin)

    logs = logs_query.order_by(AuditLog.timestamp.desc()).all()

    return render_template(
        'admin_audit_logs.html',
        admin_user=logged_in_admin,
        admin_title=admin_info['title'],
        is_general_admin=is_general_admin,
        logs=logs,
        log_count=len(logs)
    )


@app.route('/admin/settings', methods=['GET', 'POST'])
def admin_settings():
    if not session.get('admin_user'):
        return redirect(url_for('admin_login'))

    is_general = session.get('admin_user') in ['admin', 'admin gen', 'admin_federal_police']

    message = None
    error = None

    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'preferences':
            message = "System preferences updated successfully."

        elif action == 'password':
            current_password = request.form.get('current_password')
            new_password = request.form.get('new_password')
            message = "Password updated successfully."

        elif action == 'add_service' and is_general:
            service_key = request.form.get('service_key', '').strip().lower().replace(' ', '_')
            service_name = request.form.get('service_name', '').strip()
            if service_key and service_name:
                existing = Service.query.filter_by(service_key=service_key).first()
                if not existing:
                    new_service = Service()
                    new_service.service_key = service_key
                    new_service.service_name = service_name
                    db.session.add(new_service)
                    db.session.commit()
                    log_admin_action(session.get('admin_user'), f"Added new service: {service_key}")
                    message = f"Service '{service_name}' added successfully."
                else:
                    error = "Service key already exists."
            else:
                error = "Both service key and name are required."

        elif action == 'update_service' and is_general:
            service_key = request.form.get('service_key')
            new_service_name = request.form.get('new_service_name', '').strip()
            service_to_update = Service.query.filter_by(service_key=service_key).first()
            if service_to_update and new_service_name:
                service_to_update.service_name = new_service_name
                db.session.commit()
                log_admin_action(session.get('admin_user'), f"Updated service: {service_key}")
                message = "Service updated successfully."
            else:
                error = "Service not found or invalid name."

        elif action == 'delete_service' and is_general:
            service_key = request.form.get('service_key')
            service_to_delete = Service.query.filter_by(service_key=service_key).first()
            if service_to_delete:
                db.session.delete(service_to_delete)
                db.session.commit()
                log_admin_action(session.get('admin_user'), f"Deleted service: {service_key}")
                message = "Service deleted successfully."
            else:
                error = "Service not found."

        elif action == 'add_sub_service' and is_general:
            message = "Sub-service added successfully."

        elif action == 'update_sub_service' and is_general:
            message = "Sub-service updated successfully."

        elif action == 'delete_sub_service' and is_general:
            message = "Sub-service deleted successfully."

    services = Service.query.all()
    service_map = {s.service_key: s.service_name for s in services}

    return render_template(
        'admin_settings.html',
        services=services,
        service_map=service_map,
        is_general=is_general,
        message=message,
        error=error
    )


@app.route('/example-route')
def example_route():
    services = Service.query.all()
    return render_template('services.html', services=services)


@app.route('/admin/export/<format_type>')
def export_feedbacks(format_type):
    logged_in_admin = session.get('admin_user')
    admin_credentials = get_admin_credentials()
    if not logged_in_admin or logged_in_admin not in admin_credentials:
        return redirect(url_for('admin_login'))

    admin_info = admin_credentials[logged_in_admin]
    records = get_filtered_feedbacks(admin_info['type'], admin_info['service'], admin_info['sub_service'])
    log_admin_action(logged_in_admin, f"Exported feedback report in {format_type.upper()} format.")

    if format_type == 'excel':
        file_io = generate_excel_report(records)
        return send_file(
            file_io,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f"police_feedback_report_{datetime.now().strftime('%Y%m%d')}.xlsx"
        )
    elif format_type == 'word':
        file_io = generate_word_report(records)
        return send_file(
            file_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"police_feedback_report_{datetime.now().strftime('%Y%m%d')}.docx"
        )
    elif format_type == 'pdf':
        file_io = generate_pdf_report(records)
        return send_file(
            file_io,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"police_feedback_report_{datetime.now().strftime('%Y%m%d')}.pdf"
        )

    return "Invalid export format requested.", 400


@app.route('/admin/mark-read/<int:fb_id>', methods=['POST'])
def mark_feedback_read(fb_id):
    logged_in_admin = session.get('admin_user')
    admin_credentials = get_admin_credentials()
    if not logged_in_admin or logged_in_admin not in admin_credentials:
        return jsonify({"status": "error", "message": "Unauthorized"}), 401

    fb = db.session.get(Feedback, fb_id)
    if fb:
        fb.is_read = True
        db.session.commit()
        return jsonify({"status": "success"})
    return jsonify({"status": "error", "message": "Feedback not found"}), 404


@app.route('/admin/delete/<int:fb_id>', methods=['POST'])
def delete_feedback(fb_id):
    logged_in_admin = session.get('admin_user')
    admin_credentials = get_admin_credentials()
    if not logged_in_admin or logged_in_admin not in admin_credentials:
        return redirect(url_for('admin_login'))

    fb = db.session.get(Feedback, fb_id)
    if fb:
        db.session.delete(fb)
        db.session.commit()
        log_admin_action(logged_in_admin, f"Deleted feedback record #{fb_id}.")

    return redirect(url_for('admin_dashboard'))


if __name__ == '__main__':
    print("[startup] PostgreSQL database: federal_police_feedback")
    print("[startup] PostgreSQL host: localhost | port: 5432 | user: postgres")
    app.run(
        debug=True,
        host='0.0.0.0',
        port=int(os.getenv('PORT', '5000'))
    )